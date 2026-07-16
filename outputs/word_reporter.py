# -*- coding: utf-8 -*-
"""
Word 输出模块
生成活动策划案文档
"""
import os
from datetime import datetime
from typing import List, Dict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils.file_manager import get_file_manager
from utils.logger import logger


class WordReporter:
    """生成 Word 活动策划案"""
    
    def __init__(self):
        self.fm = get_file_manager()
    
    def _set_cell_border(self, cell, **kwargs):
        """设置单元格边框"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_el = tcPr.find(qn(f'w:{edge}'))
            if edge_el is None:
                edge_el = OxmlElement(f'w:{edge}')
                tcPr.append(edge_el)
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
    
    def create_activity_plan(self,
                           topics: List[Dict],
                           contents: List[Dict],
                           briefs: List[Dict],
                           date_str: str) -> str:
        """生成活动策划案 Word 文档"""
        doc = Document()
        
        # 设置中文字体
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 标题
        title = doc.add_heading(f'每日热点内容营销活动策划案', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.runs[0]
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        
        # 日期
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f'日期: {date_str}')
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        doc.add_paragraph()
        
        # 一、执行摘要
        doc.add_heading('一、执行摘要', level=1)
        summary = doc.add_paragraph()
        summary.add_run(f'本方案基于 {date_str} 各平台热点数据，结合当前营销节点，'
                       f'筛选出 {len(topics)} 个高关联热点，制定跨平台内容营销策略。'
                       f'覆盖小红书、抖音、电商站内三大渠道，预计产出内容 {len(contents)} 份。')
        
        # 二、营销节点
        doc.add_heading('二、当前营销节点', level=1)
        event = topics[0].get("marketing_event") if topics else None
        if event:
            doc.add_paragraph(f'节点名称: {event["name"]}', style='List Bullet')
            doc.add_paragraph(f'时间范围: {event["period"]}', style='List Bullet')
            doc.add_paragraph(f'主推品类: {", ".join(event["categories"])}', style='List Bullet')
            doc.add_paragraph(f'主题方向: {", ".join(event["themes"])}', style='List Bullet')
        else:
            doc.add_paragraph('当前无进行中的大型营销节点，按日常热点借势策略执行。')
        
        # 三、热点分析
        doc.add_heading('三、今日热点分析', level=1)
        for idx, topic in enumerate(topics[:5], 1):
            hot = topic["hot_topic"]
            doc.add_paragraph(f'{idx}. {hot["title"]}', style='List Number')
            p = doc.add_paragraph()
            p.add_run(f'   匹配度: {topic["match_score"]}/100 | 关联理由: {topic["match_reason"]}')
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        # 四、内容策略
        doc.add_heading('四、内容策略与执行计划', level=1)
        for content in contents:
            if content:
                doc.add_heading(f'{content["platform"]} - {content["format"]}', level=2)
                doc.add_paragraph(content["content"])
                doc.add_paragraph()
        
        # 五、达人合作
        doc.add_heading('五、达人合作 Brief', level=1)
        for brief in briefs:
            if brief:
                doc.add_heading(f'{brief["platform"]} - {brief["content_type"]} ({brief["kol_tier"]})', level=2)
                doc.add_paragraph(brief["brief"])
                doc.add_paragraph()
        
        # 六、预算预估（占位）
        doc.add_heading('六、预算与资源预估', level=1)
        doc.add_paragraph('达人合作费用: 待填写', style='List Bullet')
        doc.add_paragraph('内容制作费用: 待填写', style='List Bullet')
        doc.add_paragraph('投放推广费用: 待填写', style='List Bullet')
        
        # 七、效果指标
        doc.add_heading('七、效果追踪指标', level=1)
        doc.add_paragraph('曝光量: 目标 10万+', style='List Bullet')
        doc.add_paragraph('互动率: 目标 5%+', style='List Bullet')
        doc.add_paragraph('引流UV: 目标 5000+', style='List Bullet')
        doc.add_paragraph('转化率: 目标 2%+', style='List Bullet')
        
        # 保存
        filepath = self.fm.get_filename(date_str, "活动策划案", ".docx")
        doc.save(filepath)
        logger.info(f"[Word] 活动策划案已生成: {filepath}")
        return filepath


# 单例
_word_reporter = None

def get_word_reporter() -> WordReporter:
    global _word_reporter
    if _word_reporter is None:
        _word_reporter = WordReporter()
    return _word_reporter
